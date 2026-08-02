"""
DOCX generation service.
Creates editable Word documents with translated content.
Supports Japanese Unicode characters.

Enhanced with:
- Japanese certificate template layout
- Professional formatting with seal, photo box placeholders
- Proper field positioning for NEB Character Certificates
- Support for other document type templates
"""
import io
import logging
import os
from django.conf import settings
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)


def generate_translated_docx(document):
    """
    Generate a DOCX file using appropriate strategy per document type.
    
    For Character Certificates:
      Uses TEMPLATE FILLING — opens the approved Japanese DOCX template,
      replaces {{PLACEHOLDER}} markers with extracted field values,
      and preserves all formatting.  This is NOT paragraph translation.
    
    For other document types:
      Uses the default paragraph-based generation.

    Args:
        document: Document model instance with extracted_text (and optionally translated_text)

    Returns:
        ContentFile with the DOCX data, or None on error
    """
    # Check document type
    doc_type = document.document_type.name if document.document_type else ''

    # ─── Character Certificate: Template Filling ───
    if 'Character Certificate' in doc_type or 'character' in doc_type.lower():
        try:
            from translation.services.template_filler import generate_character_certificate_from_template
            result = generate_character_certificate_from_template(document)
            if result:
                logger.info("Character Certificate generated via template filling")
                return result
            else:
                logger.warning("Template filler returned None, falling back to legacy generation")
        except Exception as e:
            logger.error(f"Template filler failed: {e}", exc_info=True)
            logger.warning("Falling back to legacy Character Certificate generation")

    # ─── Other document types: Default generation ───
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        if not document.translated_text:
            logger.error("No translated text available for default DOCX generation")
            return None

        doc = DocxDocument()

        # Set up page margins
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # Set default font to support Japanese
        style = doc.styles['Normal']
        font = style.font
        font.name = 'MS Gothic'
        font.size = Pt(11)

        return _generate_default_docx(document, doc)

    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        logger.error(f"DOCX generation error: {e}")
        return None


def _generate_character_certificate_docx(document, doc):
    """
    Generate DOCX using Japanese NEB Character Certificate template layout.
    
    ENHANCED: Uses CharacterCertificateParser for robust field extraction
    instead of doing its own regex matching.
    """
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ENHANCED: Use parser for field extraction
    from translation.services.parsers.registry import get_parser_for_document
    
    extracted = document.extracted_text or ''
    translated = document.translated_text or ''
    
    logger.info(f"DOCX Generation - Extracted text length: {len(extracted)}")
    logger.info(f"DOCX Generation - Translated text length: {len(translated)}")
    
    # Use CharacterCertificateParser for robust extraction
    parser = get_parser_for_document('Character Certificate')
    parse_result = parser.parse(extracted)
    
    fields = parse_result.fields
    confidence = parse_result.metadata.get('confidence', 0.0)
    
    logger.info(f"Parser extracted {len(fields)} fields with {confidence:.0%} confidence")
    logger.info(f"Extracted fields: {fields}")
    
    # Get fields from parser result (with fallback to empty string)
    serial_no = fields.get('serial_no', '')
    reg_no = fields.get('reg_no', '')
    student_name = fields.get('student_name', '')
    school_name = fields.get('school_name', '')
    location = fields.get('location', '')
    year_bs = fields.get('year_bs', '')
    year_ad = fields.get('year_ad', '')
    gpa = fields.get('gpa', '')
    grade = fields.get('grade', '')
    issue_date = fields.get('issue_date', '')
    
    # Log extracted values for debugging
    logger.info(f"Serial: {serial_no}, Reg: {reg_no}")
    logger.info(f"Student: {student_name}, School: {school_name}, Location: {location}")
    logger.info(f"Grade: {grade}, GPA: {gpa}, Year BS: {year_bs}, Year AD: {year_ad}")

    # ─── HEADER: Serial No and Registration No ───
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(7)
    header_table.columns[1].width = Cm(7)

    # Serial Number (left)
    cell_serial = header_table.cell(0, 0)
    cell_serial.text = f'シリアル番号：{serial_no}' if serial_no else 'シリアル番号：'
    cell_serial.paragraphs[0].runs[0].font.size = Pt(10)
    cell_serial.paragraphs[0].runs[0].font.name = 'MS Gothic'

    # Registration Number (right)
    cell_reg = header_table.cell(0, 1)
    cell_reg.text = f'登録番号：{reg_no}' if reg_no else '登録番号：'
    cell_reg.paragraphs[0].runs[0].font.size = Pt(10)
    cell_reg.paragraphs[0].runs[0].font.name = 'MS Gothic'
    cell_reg.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                border_elem = OxmlElement(f'w:{border}')
                border_elem.set(qn('w:val'), 'none')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)

    doc.add_paragraph()  # Spacer

    # ── GOVERNMENT HEADER ───
    gov_para = doc.add_paragraph()
    gov_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gov_run = gov_para.add_run('ネパール政府')
    gov_run.font.size = Pt(14)
    gov_run.font.name = 'MS Gothic'
    gov_run.font.bold = True

    doc.add_paragraph()  # Spacer

    # ─── PHOTO & SEAL TABLE ───
    photo_table = doc.add_table(rows=1, cols=2)
    photo_table.autofit = False
    photo_table.columns[0].width = Cm(8)
    photo_table.columns[1].width = Cm(5)

    # Left side: Empty for spacing
    cell_left = photo_table.cell(0, 0)
    cell_left.text = ''

    # Right side: Photo and Seal placeholders
    cell_photo = photo_table.cell(0, 1)
    photo_para = cell_photo.paragraphs[0]
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_run = photo_para.add_run('写真\n\n印')
    photo_run.font.size = Pt(10)
    photo_run.font.name = 'MS Gothic'

    # Remove table borders
    for row in photo_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                border_elem = OxmlElement(f'w:{border}')
                border_elem.set(qn('w:val'), 'none')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # ─── CERTIFICATE TITLE ───
    cert_para = doc.add_paragraph()
    cert_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert_run = cert_para.add_run('証明書')
    cert_run.font.size = Pt(18)
    cert_run.font.name = 'MS Gothic'
    cert_run.font.bold = True
    cert_run.underline = True

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # ─── MAIN CERTIFICATION TEXT ───
    # Build certification text using extracted fields (name, school, GPA in English)
    if student_name and school_name:
        # Build school line with location
        school_line = school_name
        if location:
            school_line = f"{school_name}, {location}"
        
        # Match exact format from sample: これは、{school} の{student} 氏が、
        # Note: Single space before and after の and 氏が
        cert_text = f"これは、{school_line} の{student_name} 氏が、"
        if year_bs and year_ad:
            cert_text += f"ネパール暦{year_bs}年（西暦{year_ad}年）に"
        cert_text += f"国家試験委員会によって実施された卒業証明書試験（グレード{grade}）を"
        if gpa:
            cert_text += f"{gpa}GPAで"
        cert_text += "卒業したことを証明するものです。"  # Use 卒業 as per sample
        
        main_para = doc.add_paragraph()
        main_para.paragraph_format.line_spacing = 1.5
        main_run = main_para.add_run(cert_text)
        main_run.font.size = Pt(12)
        main_run.font.name = 'MS Gothic'
        logger.info(f"Generated certification text with: {student_name}, {school_line}, {gpa}GPA, Grade {grade}")
        logger.info(f"Cert text: {cert_text}")
    elif translated:
        # Fallback to using translated text directly
        main_text = translated.split('\n')[0]
        if main_text:
            main_para = doc.add_paragraph()
            main_para.paragraph_format.line_spacing = 1.5
            main_run = main_para.add_run(main_text)
            main_run.font.size = Pt(12)
            main_run.font.name = 'MS Gothic'

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # ─── SIGNATURE SECTION ───
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_run = sig_para.add_run('署名\n\n委員長')
    sig_run.font.size = Pt(11)
    sig_run.font.name = 'MS Gothic'

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # ─── ISSUE DATE ───
    if issue_date:
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'発行日：{issue_date}')
        date_run.font.size = Pt(10)
        date_run.font.name = 'MS Gothic'

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # ─── TRANSLATION AGENCY FOOTER ───
    # Add separator line
    sep_para = doc.add_paragraph()
    sep_para.paragraph_format.space_before = Pt(12)
    sep_run = sep_para.add_run('_' * 80)
    sep_run.font.size = Pt(8)
    sep_run.font.name = 'MS Gothic'
    sep_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # Spacer

  

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create filename
    safe_title = ''.join(c for c in document.title if c.isalnum() or c in ' _-')[:50]
    filename = f"{safe_title}_translated.docx"

    return ContentFile(buffer.read(), name=filename)


def _generate_default_docx(document, doc):
    """
    Generate default DOCX for non-certificate documents.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add title
    title = doc.add_heading(level=1)
    title_run = title.add_run(document.title)
    title_run.font.name = 'MS Gothic'

    # Add document type if available
    if document.document_type:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'Document Type: {document.document_type.name}')
        run.font.size = Pt(9)
        run.font.name = 'MS Gothic'

    doc.add_paragraph()  # Spacer

    # Add translated content paragraph by paragraph
    translated_text = document.translated_text or ''
    paragraphs = translated_text.split('\n')

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph()
            continue

        # Detect headings (simple heuristic: short lines in ALL CAPS or ending with colon)
        if (len(para_text) < 100 and para_text.isupper()) or \
           (len(para_text) < 80 and para_text.endswith(':')):
            heading = doc.add_heading(level=2)
            run = heading.add_run(para_text)
            run.font.name = 'MS Gothic'
        else:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.name = 'MS Gothic'
            run.font.size = Pt(11)

    # Add footer with metadata
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run(
        f'Translated on: {document.updated_at.strftime("%Y-%m-%d %H:%M")}'
    )
    footer_run.font.size = Pt(8)
    footer_run.font.name = 'MS Gothic'
    footer_run.font.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create filename
    safe_title = ''.join(c for c in document.title if c.isalnum() or c in ' _-')[:50]
    filename = f"{safe_title}_translated.docx"

    return ContentFile(buffer.read(), name=filename)
