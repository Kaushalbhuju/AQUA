"""
Template Filler Engine for Character Certificate DOCX generation.

This module implements "template filling" — NOT "paragraph translation".

Workflow:
  1. Open the approved Japanese DOCX template
  2. Find {{PLACEHOLDER}} markers in all paragraphs, tables, headers, footers
  3. Replace placeholders with extracted field values
  4. Preserve ALL formatting (fonts, sizes, spacing, layout, tables)
  5. Save as a new DOCX file

The generated DOCX is visually identical to the approved template,
with only student-specific information replaced.
"""
import io
import copy
import logging
import os
import re

from django.conf import settings
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)

# Path to the approved Japanese DOCX template
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docx_templates')
CHARACTER_CERTIFICATE_TEMPLATE = os.path.join(TEMPLATE_DIR, 'character_certificate_ja.docx')


# ─── Placeholder definitions ────────────────────────────────────────────────
PLACEHOLDERS = {
    '{{SERIAL_NO}}':        'serial_no',
    '{{REGISTRATION_NO}}':  'reg_no',
    '{{STUDENT_NAME}}':     'student_name',
    '{{SCHOOL_NAME}}':      'school_name',
    '{{SCHOOL_LOCATION}}':  'school_location',
    '{{GRADE}}':            'grade',
    '{{GPA}}':              'gpa',
    '{{EXAM_YEAR_BS}}':     'exam_year_bs',
    '{{EXAM_YEAR_AD}}':     'exam_year_ad',
    '{{ISSUE_DATE_BS}}':    'issue_date_bs',
    '{{ISSUE_DATE_AD}}':    'issue_date_ad',
}


def _replace_in_runs(paragraph, field_values):
    """
    Replace placeholders in a paragraph's runs while preserving formatting.
    
    Handles cases where a placeholder is split across multiple runs
    (e.g., run1="{{STUDENT", run2="_NAME}}").
    
    Strategy:
    1. First, try direct replacement in individual runs (fast path)
    2. If placeholders span multiple runs, reconstruct and re-split
    """
    # Fast path: check if any placeholder exists in individual runs
    for run in paragraph.runs:
        for placeholder, field_key in PLACEHOLDERS.items():
            if placeholder in run.text:
                value = field_values.get(field_key, '')
                run.text = run.text.replace(placeholder, value)
    
    # Check if there are still unresolved placeholders across runs
    full_text = ''.join(run.text for run in paragraph.runs)
    remaining_placeholders = [p for p in PLACEHOLDERS.keys() if p in full_text]
    
    if not remaining_placeholders:
        return  # All resolved
    
    # Slow path: handle split placeholders
    # Build a mapping of character positions to runs
    for placeholder in remaining_placeholders:
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder not in full_text:
            continue
        
        value = field_values.get(PLACEHOLDERS[placeholder], '')
        
        # Find the placeholder in the combined text
        start_idx = full_text.index(placeholder)
        end_idx = start_idx + len(placeholder)
        
        # Map positions to runs
        current_pos = 0
        start_run_idx = None
        end_run_idx = None
        
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if start_run_idx is None and run_start <= start_idx < run_end:
                start_run_idx = i
            if run_start < end_idx <= run_end:
                end_run_idx = i
                break
            
            current_pos = run_end
        
        if start_run_idx is None or end_run_idx is None:
            continue
        
        if start_run_idx == end_run_idx:
            # Placeholder is within a single run
            run = paragraph.runs[start_run_idx]
            run.text = run.text.replace(placeholder, value)
        else:
            # Placeholder spans multiple runs
            # Put the replacement value in the first run, clear the rest
            current_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                
                if i == start_run_idx:
                    # Replace from placeholder start to end of this run
                    before = run.text[:start_idx - run_start]
                    run.text = before + value
                elif start_run_idx < i < end_run_idx:
                    # Middle runs: clear completely
                    run.text = ''
                elif i == end_run_idx:
                    # Last run: remove the remainder of the placeholder
                    after_offset = end_idx - run_start
                    run.text = run.text[after_offset:]
                
                current_pos = run_end


def _replace_in_table(table, field_values):
    """Replace placeholders in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_runs(paragraph, field_values)
            # Handle nested tables
            for nested_table in cell.tables:
                _replace_in_table(nested_table, field_values)


def fill_character_certificate(fields, template_path=None):
    """
    Fill the Japanese Character Certificate template with extracted fields.
    
    This is a "template filling" operation, NOT "paragraph translation".
    
    Args:
        fields (dict): Extracted fields from CharacterCertificateParser.
            Expected keys:
                - serial_no: e.g., "C0034274"
                - reg_no: e.g., "845271150233" 
                - student_name: e.g., "SUNIL B.K." (kept in English)
                - school_name: e.g., "NAVODIT VIDYA KUNJA SECONDARY SCHOOL" (kept in English)
                - school_location: e.g., "SAMAKHUSHI, KATHMANDU" (kept in English)
                - grade: e.g., "XII"
                - gpa: e.g., "3.14"
                - exam_year_bs: e.g., "2082"
                - exam_year_ad: e.g., "2025"
                - issue_date_bs: e.g., "2082/12/18"
                - issue_date_ad: e.g., "2026/04/01"
        
        template_path (str, optional): Path to template DOCX file.
            Defaults to the approved Japanese template.
    
    Returns:
        ContentFile with the filled DOCX, or None on error.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return None
    
    if template_path is None:
        template_path = CHARACTER_CERTIFICATE_TEMPLATE
    
    if not os.path.exists(template_path):
        logger.error(f"Template not found: {template_path}")
        logger.info("Attempting to regenerate template...")
        try:
            from translation.docx_templates.create_template import create_character_certificate_template
            create_character_certificate_template()
        except Exception as e:
            logger.error(f"Failed to regenerate template: {e}")
            return None
    
    logger.info(f"Opening template: {template_path}")
    logger.info(f"Fields to fill: {fields}")
    
    # Map parser field names to template placeholder field names
    field_values = _map_parser_fields_to_template(fields)
    
    logger.info(f"Mapped field values: {field_values}")
    
    try:
        # Open the template document
        doc = DocxDocument(template_path)
        
        # Replace placeholders in all paragraphs (main body)
        for paragraph in doc.paragraphs:
            _replace_in_runs(paragraph, field_values)
        
        # Replace placeholders in all tables
        for table in doc.tables:
            _replace_in_table(table, field_values)
        
        # Replace in headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    _replace_in_runs(paragraph, field_values)
                for table in section.header.tables:
                    _replace_in_table(table, field_values)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    _replace_in_runs(paragraph, field_values)
                for table in section.footer.tables:
                    _replace_in_table(table, field_values)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create filename from student name
        student_name = field_values.get('student_name', 'unknown')
        safe_name = ''.join(c for c in student_name if c.isalnum() or c in ' _-')[:40]
        filename = f"Character_Certificate_{safe_name}_JP.docx"
        
        logger.info(f"Template filled successfully: {filename}")
        return ContentFile(buffer.read(), name=filename)
        
    except Exception as e:
        logger.error(f"Template filling failed: {e}", exc_info=True)
        return None


def _map_parser_fields_to_template(fields):
    """
    Map fields from CharacterCertificateParser output to template placeholder names.
    
    The parser outputs fields like 'serial_no', 'reg_no', 'year_bs', etc.
    The template expects 'serial_no', 'reg_no', 'exam_year_bs', etc.
    
    This function handles the mapping and also splits combined fields
    (e.g., issue_date "2082/12/18 (4/1/2026)" into issue_date_bs and issue_date_ad).
    """
    mapped = {}
    
    # Direct mappings
    mapped['serial_no'] = fields.get('serial_no', '')
    mapped['reg_no'] = fields.get('reg_no', '')
    mapped['student_name'] = fields.get('student_name', '')
    mapped['school_name'] = fields.get('school_name', '')
    mapped['school_location'] = fields.get('location', fields.get('school_location', ''))
    mapped['grade'] = fields.get('grade', '')
    mapped['gpa'] = fields.get('gpa', '')
    
    # Year mappings
    mapped['exam_year_bs'] = fields.get('year_bs', fields.get('exam_year_bs', ''))
    mapped['exam_year_ad'] = fields.get('year_ad', fields.get('exam_year_ad', ''))
    
    # Issue date - may need splitting
    issue_date = fields.get('issue_date', '')
    issue_date_bs = fields.get('issue_date_bs', '')
    issue_date_ad = fields.get('issue_date_ad', '')
    
    if issue_date and not issue_date_bs:
        # Try to parse combined format: "2082/12/18 (4/1/2026)"
        bs_match = re.search(r'(\d{4}/\d{1,2}/\d{1,2})', issue_date)
        ad_match = re.search(r'\((\d{1,2}/\d{1,2}/\d{4})\)', issue_date)
        
        if bs_match:
            issue_date_bs = bs_match.group(1)
        if ad_match:
            # Convert from M/D/YYYY to YYYY/MM/DD
            ad_raw = ad_match.group(1)
            ad_parts = ad_raw.split('/')
            if len(ad_parts) == 3:
                issue_date_ad = f"{ad_parts[2]}/{ad_parts[0].zfill(2)}/{ad_parts[1].zfill(2)}"
            else:
                issue_date_ad = ad_raw
        
        # If only BS date found, use it as-is
        if bs_match and not ad_match:
            issue_date_bs = bs_match.group(1)
        
        # If no pattern matched, use the raw issue_date as BS
        if not issue_date_bs:
            issue_date_bs = issue_date
    
    mapped['issue_date_bs'] = issue_date_bs
    mapped['issue_date_ad'] = issue_date_ad
    
    return mapped


def generate_character_certificate_from_template(document):
    """
    Generate a Character Certificate DOCX by filling the Japanese template.
    
    This is the main entry point called from docx_generator.py.
    
    Instead of translating paragraphs, this:
    1. Parses the extracted English text to get fields
    2. Opens the approved Japanese template
    3. Fills placeholders with extracted values
    4. Returns the filled DOCX
    
    Args:
        document: Document model instance with extracted_text
    
    Returns:
        ContentFile with filled DOCX, or None on error
    """
    from translation.services.parsers.registry import get_parser_for_document
    
    extracted = document.extracted_text or ''
    
    if not extracted:
        logger.error("No extracted text available for template filling")
        return None
    
    # Parse extracted text to get structured fields
    parser = get_parser_for_document('Character Certificate')
    parse_result = parser.parse(extracted)
    
    fields = parse_result.fields
    confidence = parse_result.metadata.get('confidence', 0.0)
    
    logger.info(f"Template filler: {len(fields)} fields extracted with {confidence:.0%} confidence")
    logger.info(f"Fields: {fields}")
    
    # Fill the template
    return fill_character_certificate(fields)
