"""
Script to generate the approved Japanese Character Certificate DOCX template.
This template contains fixed Japanese text with {{PLACEHOLDER}} markers
that will be replaced with extracted field values.

Run this once to generate the template:
  python translation/docx_templates/create_template.py
"""
import os
import sys

# Add project root to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def remove_table_borders(table):
    """Remove all borders from a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:val'), 'none')
                border_elem.set(qn('w:sz'), '0')
                border_elem.set(qn('w:space'), '0')
                border_elem.set(qn('w:color'), 'auto')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)


def add_cell_border(cell, borders_dict):
    """Add specific borders to a cell.
    borders_dict: {'top': {'val': 'single', 'sz': '4', 'color': '000000'}, ...}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, props in borders_dict.items():
        border_elem = OxmlElement(f'w:{border_name}')
        for prop_name, prop_val in props.items():
            border_elem.set(qn(f'w:{prop_name}'), prop_val)
        tcBorders.append(border_elem)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)


def add_run(paragraph, text, font_name='MS Gothic', font_size=11, bold=False, 
            underline=False, color=None):
    """Add a run with specified formatting."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    
    # Ensure East Asian font fallback
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    
    return run


def create_character_certificate_template():
    """
    Create the approved Japanese Character Certificate template
    matching the consultancy's approved layout.
    
    Placeholders:
      {{SERIAL_NO}}       - Serial number (e.g., C0034274)
      {{REGISTRATION_NO}} - Registration number (e.g., 845271150233)
      {{STUDENT_NAME}}    - Student name in English (e.g., SUNIL B.K.)
      {{SCHOOL_NAME}}     - School name in English
      {{SCHOOL_LOCATION}} - School location in English
      {{GRADE}}           - Grade (XI or XII)
      {{GPA}}             - GPA (e.g., 3.14)
      {{EXAM_YEAR_BS}}    - Exam year in B.S. (e.g., 2082)
      {{EXAM_YEAR_AD}}    - Exam year in A.D. (e.g., 2025)
      {{ISSUE_DATE_BS}}   - Issue date in B.S. (e.g., 2082/12/18)
      {{ISSUE_DATE_AD}}   - Issue date in A.D. (e.g., 2026/04/01)
    """
    doc = Document()
    
    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_width = Cm(21)      # A4
    section.page_height = Cm(29.7)   # A4
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'MS Gothic'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    
    # ─── ROW 1: Serial No (left) and Registration No (right) ───
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(8)
    header_table.columns[1].width = Cm(8)
    
    # Serial Number (left)
    cell_serial = header_table.cell(0, 0)
    p_serial = cell_serial.paragraphs[0]
    p_serial.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p_serial, 'シリアル番号：', font_size=10)
    add_run(p_serial, '{{SERIAL_NO}}', font_size=10)
    
    # Registration Number (right)
    cell_reg = header_table.cell(0, 1)
    p_reg = cell_reg.paragraphs[0]
    p_reg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_reg, '登録番号：', font_size=10)
    add_run(p_reg, '{{REGISTRATION_NO}}', font_size=10)
    
    remove_table_borders(header_table)
    
    # ─── Spacer ───
    doc.add_paragraph()
    
    # ─── GOVERNMENT HEADER ───
    gov_para = doc.add_paragraph()
    gov_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(gov_para, 'ネパール政府', font_size=14, bold=True)
    
    # ─── National Examinations Board (large) ───
    # Use a table to place シル (seal) on the left and 国家試験委員会 centered
    neb_table = doc.add_table(rows=1, cols=3)
    neb_table.autofit = False
    neb_table.columns[0].width = Cm(3)
    neb_table.columns[1].width = Cm(10)
    neb_table.columns[2].width = Cm(3)
    
    # Left: Seal circle placeholder (シル)
    cell_seal_left = neb_table.cell(0, 0)
    p_seal = cell_seal_left.paragraphs[0]
    p_seal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_seal, 'シル', font_size=10)
    
    # Center: 国家試験委員会
    cell_neb = neb_table.cell(0, 1)
    p_neb = cell_neb.paragraphs[0]
    p_neb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_neb, '国家試験委員会', font_size=22, bold=True)
    
    # Right: empty
    neb_table.cell(0, 2).paragraphs[0].text = ''
    
    remove_table_borders(neb_table)
    
    # ─── Stamp and Photo area ───
    stamp_table = doc.add_table(rows=1, cols=3)
    stamp_table.autofit = False
    stamp_table.columns[0].width = Cm(5)
    stamp_table.columns[1].width = Cm(5)
    stamp_table.columns[2].width = Cm(6)
    
    # Left: empty
    stamp_table.cell(0, 0).paragraphs[0].text = ''
    
    # Center: 印 (stamp) placeholder
    cell_stamp = stamp_table.cell(0, 1)
    p_stamp = cell_stamp.paragraphs[0]
    p_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_stamp, '印', font_size=10)
    
    # Right: 写真 (photo) placeholder with border
    cell_photo = stamp_table.cell(0, 2)
    p_photo = cell_photo.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_photo, '\n写真\n\n', font_size=10)
    add_cell_border(cell_photo, {
        'top': {'val': 'single', 'sz': '4', 'color': '000000'},
        'bottom': {'val': 'single', 'sz': '4', 'color': '000000'},
        'left': {'val': 'single', 'sz': '4', 'color': '000000'},
        'right': {'val': 'single', 'sz': '4', 'color': '000000'},
    })
    
    remove_table_borders(stamp_table)
    # Re-add border for photo cell only
    add_cell_border(cell_photo, {
        'top': {'val': 'single', 'sz': '4', 'color': '000000'},
        'bottom': {'val': 'single', 'sz': '4', 'color': '000000'},
        'left': {'val': 'single', 'sz': '4', 'color': '000000'},
        'right': {'val': 'single', 'sz': '4', 'color': '000000'},
    })
    
    # ─── Spacer ───
    doc.add_paragraph()
    
    # ─── CERTIFICATE TITLE ───
    cert_para = doc.add_paragraph()
    cert_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cert_para, '証明書', font_size=18, bold=True, underline=True)
    
    # ─── Spacer ───
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ─── MAIN CERTIFICATION TEXT ───
    # This is the core template paragraph with placeholders
    main_para = doc.add_paragraph()
    main_para.paragraph_format.line_spacing = 1.8
    
    add_run(main_para, 'これは、', font_size=12)
    add_run(main_para, '{{SCHOOL_NAME}}', font_size=12)
    add_run(main_para, '、', font_size=12)
    add_run(main_para, '{{SCHOOL_LOCATION}}', font_size=12)
    add_run(main_para, ' の', font_size=12)
    add_run(main_para, '{{STUDENT_NAME}}', font_size=12)
    add_run(main_para, ' 氏が、ネパール暦', font_size=12)
    add_run(main_para, '{{EXAM_YEAR_BS}}', font_size=12)
    add_run(main_para, '年（西暦', font_size=12)
    add_run(main_para, '{{EXAM_YEAR_AD}}', font_size=12)
    add_run(main_para, '年）に国家試験委員会によって実施された卒業証明書試験（グレード', font_size=12)
    add_run(main_para, '{{GRADE}}', font_size=12)
    add_run(main_para, '）を', font_size=12)
    add_run(main_para, '{{GPA}}', font_size=12)
    add_run(main_para, ' GPAで卒業したことを証明するものです。', font_size=12)
    
    # ─── Spacers ───
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ─── SIGNATURE SECTION (right-aligned) ───
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(10)
    sig_table.columns[1].width = Cm(6)
    
    # Empty left cells
    sig_table.cell(0, 0).paragraphs[0].text = ''
    sig_table.cell(1, 0).paragraphs[0].text = ''
    
    # Right: 署名 (Signature) label
    cell_sig = sig_table.cell(0, 1)
    p_sig = cell_sig.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_sig, '署名', font_size=11)
    
    # Right: 委員長 (Chairperson) label
    cell_chair = sig_table.cell(1, 1)
    p_chair = cell_chair.paragraphs[0]
    p_chair.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p_chair, '委員長', font_size=11)
    
    remove_table_borders(sig_table)
    
    # ─── Spacer ───
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ─── ISSUE DATE ───
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(date_para, '発行日：', font_size=10)
    add_run(date_para, '{{ISSUE_DATE_BS}}', font_size=10)
    add_run(date_para, '（', font_size=10)
    add_run(date_para, '{{ISSUE_DATE_AD}}', font_size=10)
    add_run(date_para, '）', font_size=10)
    
    # ─── Save template ───
    template_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(template_dir, 'character_certificate_ja.docx')
    doc.save(template_path)
    print(f"Template saved to: {template_path}")
    return template_path


if __name__ == '__main__':
    create_character_certificate_template()
